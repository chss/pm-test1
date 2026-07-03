import streamlit as st
import json
import pandas as pd
import os
import io
from io import BytesIO

# Try imports for optional dependencies
try:
    from databricks import sql
    HAS_DATABRICKS = True
except ImportError:
    HAS_DATABRICKS = False

try:
    from google import genai
    from google.genai import types
    HAS_GEMINI = True
except ImportError:
    HAS_GEMINI = False

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Set page config
st.set_page_config(
    page_title="Enterprise Business Case Intake Form",
    page_icon="💼",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Premium CSS for Light Theme Design
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;700&display=swap');
    
    /* Font styles */
    html, body, [class*="css"], .stMarkdown {
        font-family: 'Outfit', sans-serif;
    }
    
    /* Title styling */
    .main-title {
        background: linear-gradient(135deg, #4f46e5 0%, #7c3aed 50%, #db2777 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-size: 3rem;
        font-weight: 700;
        margin-bottom: 0.5rem;
        animation: fadeIn 1s ease-in-out;
    }
    
    .subtitle {
        color: #475569;
        font-size: 1.1rem;
        margin-bottom: 2rem;
    }
    
    /* Custom cards for tabs & sections */
    .glass-card {
        background: #ffffff;
        border: 1px solid #e2e8f0;
        border-radius: 16px;
        padding: 24px;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05), 0 2px 4px -1px rgba(0, 0, 0, 0.03);
        margin-bottom: 20px;
    }
    
    .glass-card h3 {
        color: #0f172a;
        margin-top: 0;
        margin-bottom: 15px;
        border-bottom: 1px solid #e2e8f0;
        padding-bottom: 8px;
    }
    
    /* Sidebar styling: White background & black text */
    [data-testid="stSidebar"] {
        background-color: #ffffff !important;
        border-right: 1px solid #e2e8f0 !important;
    }
    
    [data-testid="stSidebar"] .sidebar-section {
        background: #f8fafc !important;
        border-radius: 12px;
        padding: 15px;
        margin-bottom: 15px;
        border: 1px solid #e2e8f0 !important;
    }
    
    [data-testid="stSidebar"] h2, 
    [data-testid="stSidebar"] h3, 
    [data-testid="stSidebar"] h4, 
    [data-testid="stSidebar"] p, 
    [data-testid="stSidebar"] label,
    [data-testid="stSidebar"] span,
    [data-testid="stSidebar"] small,
    [data-testid="stSidebar"] .progress-label {
        color: #0f172a !important;
    }
    
    /* Inputs, text fields and download buttons in sidebar */
    [data-testid="stSidebar"] input {
        background-color: #ffffff !important;
        color: #0f172a !important;
        border: 1px solid #cbd5e1 !important;
    }
    
    [data-testid="stSidebar"] [data-testid="stExpander"] {
        background-color: #f8fafc !important;
        border: 1px solid #cbd5e1 !important;
        border-radius: 8px;
    }

    [data-testid="stSidebar"] [data-testid="stExpander"] * {
        background-color: #f8fafc !important;
        color: #0f172a !important;
    }
    
    /* Progress bar custom styling */
    .progress-label {
        font-size: 0.85rem;
        font-weight: 600;
        color: #1e293b !important;
        display: flex;
        justify-content: space-between;
        margin-bottom: 5px;
    }
    
    /* Key metrics indicator */
    .status-badge {
        font-size: 0.75rem;
        padding: 4px 8px;
        border-radius: 9999px;
        font-weight: 600;
    }
    
    .badge-complete {
        background-color: rgba(16, 185, 129, 0.15) !important;
        color: #059669 !important;
        border: 1px solid rgba(16, 185, 129, 0.3) !important;
    }
    
    .badge-incomplete {
        background-color: rgba(245, 158, 11, 0.15) !important;
        color: #d97706 !important;
        border: 1px solid rgba(245, 158, 11, 0.3) !important;
    }

    /* Tab styling overrides */
    .stTabs [data-baseweb="tab-list"] {
        gap: 10px;
    }

    .stTabs [data-baseweb="tab"] {
        background-color: #f1f5f9;
        border: 1px solid #e2e8f0;
        border-radius: 8px 8px 0px 0px;
        padding: 10px 20px;
        color: #475569;
        font-weight: 600;
        transition: all 0.3s ease;
    }

    .stTabs [data-baseweb="tab"]:hover {
        color: #0f172a;
        background-color: #e2e8f0;
    }

    .stTabs [aria-selected="true"] {
        background-color: #ffffff !important;
        color: #4f46e5 !important;
        border-color: #cbd5e1 #cbd5e1 #ffffff #cbd5e1 !important;
    }
    
    /* Animation definition */
    @keyframes fadeIn {
        from { opacity: 0; transform: translateY(-10px); }
        to { opacity: 1; transform: translateY(0); }
    }
</style>
""", unsafe_allow_html=True)

# Initialize Session State for inputs
if "project_name" not in st.session_state:
    st.session_state.project_name = ""
if "problem_solving" not in st.session_state:
    st.session_state.problem_solving = ""
if "who_impacted" not in st.session_state:
    st.session_state.who_impacted = ""
if "cost_impact" not in st.session_state:
    st.session_state.cost_impact = ""
if "other_details" not in st.session_state:
    st.session_state.other_details = ""

# Value creation
if "value_creation" not in st.session_state:
    st.session_state.value_creation = ""

# Dynamic KPIs list
if "kpis" not in st.session_state:
    st.session_state.kpis = [
        {
            "kpi": "",
            "alignment": "",
            "measured": "",
            "easily_measurable": ""
        }
    ]

# Business Impacts options
if "impact_revenue" not in st.session_state:
    st.session_state.impact_revenue = "No Impact"
if "impact_cost" not in st.session_state:
    st.session_state.impact_cost = "No Impact"
if "impact_cust_service" not in st.session_state:
    st.session_state.impact_cust_service = "No impact"
if "impact_efficiency" not in st.session_state:
    st.session_state.impact_efficiency = "No impact"
if "impact_duration" not in st.session_state:
    st.session_state.impact_duration = "No impact"
if "impact_quality" not in st.session_state:
    st.session_state.impact_quality = "No impact"

# Database Configuration Session State
if "db_host" not in st.session_state:
    st.session_state.db_host = os.getenv("DATABRICKS_SERVER_HOSTNAME", "")
if "db_path" not in st.session_state:
    st.session_state.db_path = os.getenv("DATABRICKS_HTTP_PATH", "")
if "db_token" not in st.session_state:
    st.session_state.db_token = os.getenv("DATABRICKS_TOKEN", "")
if "db_catalog" not in st.session_state:
    st.session_state.db_catalog = os.getenv("DATABRICKS_CATALOG", "pm_test1")
if "db_schema" not in st.session_state:
    st.session_state.db_schema = os.getenv("DATABRICKS_SCHEMA", "pm_test1_schema")
if "db_table" not in st.session_state:
    st.session_state.db_table = os.getenv("DATABRICKS_TABLE", "business_case_intake")

# Copilot configuration & chat history
if "gemini_api_key" not in st.session_state:
    st.session_state.gemini_api_key = os.getenv("GEMINI_API_KEY", "")
if "copilot_messages" not in st.session_state:
    st.session_state.copilot_messages = []

# Session state variables for Priority Agent
if "priority_run_session_id" not in st.session_state:
    st.session_state.priority_run_session_id = None
if "pending_priority_confirmation" not in st.session_state:
    st.session_state.pending_priority_confirmation = None

# Sample Business Case Data
sample_case = {
    "project_name": "Smart Inventory Optimization Engine",
    "problem_solving": "Our regional distribution centers experience periodic stockouts of high-demand items while carrying excess safety stock of low-turnover items. This is caused by outdated static reorder points that do not account for supply lead times or seasonal trends.",
    "who_impacted": "Distribution center managers, customer experience teams, finance (holding costs), and end customers facing delayed orders.",
    "cost_impact": "Estimated annual loss of $450,000 in unfulfilled orders and $180,000 in excess inventory warehousing costs.",
    "other_details": "This engine will integrate directly with our existing SAP ERP system and run daily forecast updates.",
    "value_creation": "By optimizing stock levels, we release tied-up working capital and improve delivery compliance, directly raising net promoter score (NPS) and increasing customer repeat purchase rates.",
    "kpis": [
        {
            "kpi": "Reduce average stockout rate of tier-A products to below 1.5%",
            "alignment": "Supply Chain Excellence & Customer Satisfaction Goals",
            "measured": "ERP system inventory tracking vs daily order demand",
            "easily_measurable": "Yes, standard telemetry reports generated weekly."
        },
        {
            "kpi": "Decrease inventory carrying costs by 15%",
            "alignment": "Working Capital Optimization Target",
            "measured": "Finance balance sheet inventory valuations monthly",
            "easily_measurable": "Yes, calculated directly from standard cost accounting metrics."
        }
    ],
    "impact_revenue": "Increase",
    "impact_cost": "Saving",
    "impact_cust_service": "External",
    "impact_efficiency": "Improvement",
    "impact_duration": "Reduction",
    "impact_quality": "Product"
}

# Action: Load sample data
def load_sample_data():
    st.session_state.project_name = sample_case["project_name"]
    st.session_state.problem_solving = sample_case["problem_solving"]
    st.session_state.who_impacted = sample_case["who_impacted"]
    st.session_state.cost_impact = sample_case["cost_impact"]
    st.session_state.other_details = sample_case["other_details"]
    st.session_state.value_creation = sample_case["value_creation"]
    st.session_state.kpis = [dict(k) for k in sample_case["kpis"]]
    st.session_state.impact_revenue = sample_case["impact_revenue"]
    st.session_state.impact_cost = sample_case["impact_cost"]
    st.session_state.impact_cust_service = sample_case["impact_cust_service"]
    st.session_state.impact_efficiency = sample_case["impact_efficiency"]
    st.session_state.impact_duration = sample_case["impact_duration"]
    st.session_state.impact_quality = sample_case["impact_quality"]

# Action: Reset form
def reset_form():
    st.session_state.project_name = ""
    st.session_state.problem_solving = ""
    st.session_state.who_impacted = ""
    st.session_state.cost_impact = ""
    st.session_state.other_details = ""
    st.session_state.value_creation = ""
    st.session_state.kpis = [
        {
            "kpi": "",
            "alignment": "",
            "measured": "",
            "easily_measurable": ""
        }
    ]
    st.session_state.impact_revenue = "No Impact"
    st.session_state.impact_cost = "No Impact"
    st.session_state.impact_cust_service = "No impact"
    st.session_state.impact_efficiency = "No impact"
    st.session_state.impact_duration = "No impact"
    st.session_state.impact_quality = "No impact"

# Title header block
st.markdown('<div class="main-title">Business Case Intake Portal</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Design, validate, and outline project feasibility, KPIs, and corporate impact.</div>', unsafe_allow_html=True)

# Helper function to compute field completion stats (Only Project Name and Problem Description are required)
def get_validation_stats():
    required_fields = [
        st.session_state.get("project_name", ""),
        st.session_state.get("problem_solving", "")
    ]
    
    filled = sum(1 for f in required_fields if str(f).strip())
    total = len(required_fields)
    
    percentage = int((filled / total) * 100) if total > 0 else 100
    return filled, total, percentage

# Sanitizes all optional text fields: empty values are filled with "Not Applicable"
def sanitize_empty_fields():
    # Main form text fields
    optional_fields = ["who_impacted", "cost_impact", "other_details", "value_creation"]
    for field in optional_fields:
        val = st.session_state.get(field, "")
        if not val or not str(val).strip():
            st.session_state[field] = "Not Applicable"
            
    # KPI metric list fields
    for kpi in st.session_state.get("kpis", []):
        for key in ["kpi", "alignment", "measured", "easily_measurable"]:
            val = kpi.get(key, "")
            if not val or not str(val).strip():
                kpi[key] = "Not Applicable"

# Database Submission Helper
def save_to_databricks(data):
    if not HAS_DATABRICKS:
        st.error("❌ Databricks SQL Connector package is missing. Ensure you run 'pip install databricks-sql-connector'.")
        return False
        
    host = st.session_state.db_host.strip()
    path = st.session_state.db_path.strip()
    token = st.session_state.db_token.strip()
    catalog = st.session_state.db_catalog.strip()
    schema = st.session_state.db_schema.strip()
    table = st.session_state.db_table.strip()
    
    if not host or not path:
        st.error("❌ Server Hostname and HTTP Path are required. Please set them in the sidebar.")
        return False
        
    full_table_name = f"`{catalog}`.`{schema}`.`{table}`"
    
    try:
        # Establish Connection
        if token:
            conn = sql.connect(
                server_hostname=host,
                http_path=path,
                access_token=token
            )
        else:
            # Fallback to Databricks OAuth if deployed in Databricks Apps
            try:
                from databricks.sdk.core import Config as SDKConfig
                cfg = SDKConfig()
                conn = sql.connect(
                    server_hostname=host,
                    http_path=path,
                    credentials_provider=lambda: cfg.authenticate
                )
            except Exception as e_oauth:
                st.error(f"❌ Failed to resolve Databricks OAuth credentials: {str(e_oauth)}. Please configure a Token.")
                return False
                
        with conn.cursor() as cursor:
            # Ensure catalog and schema exist
            cursor.execute(f"CREATE DATABASE IF NOT EXISTS `{catalog}`.`{schema}`")
            
            # Create Table if it does not exist
            create_table_sql = f"""
            CREATE TABLE IF NOT EXISTS {full_table_name} (
                project_name STRING,
                problem_solving STRING,
                who_impacted STRING,
                cost_impact STRING,
                other_details STRING,
                value_creation STRING,
                kpis STRING,
                impact_revenue STRING,
                impact_cost STRING,
                impact_cust_service STRING,
                impact_efficiency STRING,
                impact_duration STRING,
                impact_quality STRING,
                priority STRING,
                submitted_at TIMESTAMP
            ) USING DELTA
            """
            cursor.execute(create_table_sql)
            
            # Ensure priority column exists if table was previously created without it
            try:
                cursor.execute(f"ALTER TABLE {full_table_name} ADD COLUMN IF NOT EXISTS priority STRING")
            except Exception:
                pass
            
            # Serialize KPIs to JSON string
            kpis_json = json.dumps(data["kpis"])
            
            # Insert using parameterized SQL statement
            insert_sql = f"""
            INSERT INTO {full_table_name} VALUES (
                ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, current_timestamp()
            )
            """
            cursor.execute(insert_sql, (
                data["project_name"],
                data["problem_solving"],
                data["who_impacted"],
                data["cost_impact"],
                data["other_details"],
                data["value_creation"],
                kpis_json,
                data["impact_revenue"],
                data["impact_cost"],
                data["impact_cust_service"],
                data["impact_efficiency"],
                data["impact_duration"],
                data["impact_quality"],
                data.get("priority", "C")
            ))
            
        conn.close()
        return True
    except Exception as e:
        st.error(f"❌ Databricks Connection or SQL error: {str(e)}")
        return False

# Priority Agent Workflow Execution Helper Functions (ADK 2.0 Graph Workflow API)
async def run_priority_workflow_async(project_data, user_response=None, invocation_id=None):
    from google.adk.runners import InMemoryRunner
    from google.adk.apps import App, ResumabilityConfig
    from google.genai import types
    from priority_workflow import priority_workflow
    import uuid
    import json
    
    if "priority_runner" not in st.session_state:
        app = App(
            name="priority_app",
            root_agent=priority_workflow,
            resumability_config=ResumabilityConfig(is_resumable=True)
        )
        st.session_state.priority_runner = InMemoryRunner(app=app)
        
    runner = st.session_state.priority_runner
    
    session_id = st.session_state.get("priority_run_session_id")
    if not session_id:
        session_id = str(uuid.uuid4())
        st.session_state.priority_run_session_id = session_id
        
        # Create session only for a brand new run to preserve history on subsequent resumes
        await runner.session_service.create_session(
            app_name=runner.app_name,
            user_id="streamlit_user",
            session_id=session_id
        )
    
    if invocation_id and user_response:
        resume_msg = types.Content(role='user', parts=[
            types.Part(
                function_response=types.FunctionResponse(
                    name='adk_request_input',
                    id='confirm_priority',
                    response={'confirm_priority': user_response}
                )
            )
        ])
        generator = runner.run_async(
            user_id="streamlit_user",
            session_id=session_id,
            invocation_id=invocation_id,
            new_message=resume_msg
        )
    else:
        # Initial run
        payload_str = json.dumps(project_data)
        new_msg = types.Content(
            role="user",
            parts=[types.Part.from_text(text=payload_str)]
        )
        generator = runner.run_async(
            user_id="streamlit_user",
            session_id=session_id,
            new_message=new_msg
        )
        
    result_event = None
    async for event in generator:
        result_event = event
        # If it's an interrupt, we can break early to let the user respond
        if getattr(event, 'content', None) and event.content.parts:
            for part in event.content.parts:
                if part.function_call and part.function_call.name == "adk_request_input":
                    return event
                    
    return result_event

def run_priority_workflow_sync(project_data, user_response=None, invocation_id=None):
    import asyncio
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(run_priority_workflow_async(project_data, user_response, invocation_id))
    finally:
        loop.close()

def initiate_priority_and_save(payload):
    # Sanitize optional text fields to "Not Applicable" if empty
    sanitize_empty_fields()
    
    # Rebuild payload to use the sanitized values from session state
    payload = {
        "project_name": st.session_state.project_name,
        "problem_solving": st.session_state.problem_solving,
        "who_impacted": st.session_state.who_impacted,
        "cost_impact": st.session_state.cost_impact,
        "other_details": st.session_state.other_details,
        "value_creation": st.session_state.value_creation,
        "kpis": st.session_state.kpis,
        "impact_revenue": st.session_state.impact_revenue,
        "impact_cost": st.session_state.impact_cost,
        "impact_cust_service": st.session_state.impact_cust_service,
        "impact_efficiency": st.session_state.impact_efficiency,
        "impact_duration": st.session_state.impact_duration,
        "impact_quality": st.session_state.impact_quality
    }
    
    # Compile inputs for Priority Agent graph schema
    agent_input = {
        "project_name": payload["project_name"],
        "impact_revenue": payload["impact_revenue"],
        "impact_cost": payload["impact_cost"],
        "impact_cust_service": payload["impact_cust_service"],
        "impact_quality": payload["impact_quality"]
    }
    
    with st.spinner("Invoking Priority Agent..."):
        # Run Priority Workflow (ADK 2.0 Graph Workflow API)
        event = run_priority_workflow_sync(agent_input)
        
        is_interrupted = False
        if event and getattr(event, 'content', None) and event.content.parts:
            for part in event.content.parts:
                if part.function_call and part.function_call.name == "adk_request_input":
                    interrupt_id = part.function_call.args.get("interruptId")
                    msg = part.function_call.args.get("message")
                    calc_priority = msg.split(":")[-1]
                    
                    # Store confirmation request in session state
                    st.session_state.pending_priority_confirmation = {
                        "calculated_priority": calc_priority,
                        "invocation_id": event.invocation_id,
                        "payload": payload
                    }
                    is_interrupted = True
                    st.rerun()
                    
        if not is_interrupted:
            priority = "C"
            if event and event.output and "priority" in event.output:
                priority = event.output["priority"]
            payload["priority"] = priority
            
            success = save_to_databricks(payload)
            if success:
                st.balloons()
                st.success(f"🎉 Saved successfully to `{st.session_state.db_catalog}`.`{st.session_state.db_schema}`.`{st.session_state.db_table}`!")

# Word Document Generator Utility
def generate_project_charter_docx():
    if not HAS_DOCX:
        st.error("❌ 'python-docx' library is not available. Please add it to requirements.")
        return None
        
    # Sanitize optional text fields to "Not Applicable" if empty
    sanitize_empty_fields()
    
    doc = Document()
    
    # Custom Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Helper to color table cell background
    def set_cell_background(cell, fill_hex):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))
        
    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("PROJECT CHARTER")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(79, 70, 229) # Indigo
    
    # Project Name Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    proj_name = st.session_state.project_name if st.session_state.project_name.strip() else "Unnamed Project"
    run_sub = p_sub.add_run(proj_name.upper())
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(14)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(71, 85, 105)
    
    # Horizontal line divider
    p_div = doc.add_paragraph()
    p_div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_div.add_run("―" * 40).font.color.rgb = RGBColor(226, 232, 240)
    
    # Section 1: Problem Statement
    h1 = doc.add_heading(level=1)
    r1 = h1.add_run("1. Problem Statement & Context")
    r1.font.color.rgb = RGBColor(79, 70, 229)
    
    def add_field_section(title, content):
        p_label = doc.add_paragraph()
        run_l = p_label.add_run(f"{title}:")
        run_l.font.bold = True
        run_l.font.size = Pt(11)
        run_l.font.color.rgb = RGBColor(15, 23, 42)
        
        p_content = doc.add_paragraph()
        val = content.strip() if content.strip() else "Not specified."
        run_c = p_content.add_run(val)
        run_c.font.size = Pt(11)
        run_c.font.color.rgb = RGBColor(71, 85, 105)
        p_content.paragraph_format.left_indent = Inches(0.2)
        p_content.paragraph_format.space_after = Pt(12)
        
    add_field_section("What Specific Problem Are We Solving?", st.session_state.problem_solving)
    add_field_section("Who is Impacted?", st.session_state.who_impacted)
    add_field_section("What is the Current Cost/Impact of Not Solving the Problem?", st.session_state.cost_impact)
    add_field_section("Other Details", st.session_state.other_details)
    
    # Section 2: Value Creation & KPIs
    h2 = doc.add_heading(level=1)
    r2 = h2.add_run("2. Value Creation & Key Performance Indicators")
    r2.font.color.rgb = RGBColor(79, 70, 229)
    
    add_field_section("How the Project Directly Relates to Value Creation", st.session_state.value_creation)
    
    # Add KPIs Table
    p_kpi_lbl = doc.add_paragraph()
    run_kpi_lbl = p_kpi_lbl.add_run("Key Performance Indicators:")
    run_kpi_lbl.font.bold = True
    
    kpis_table = doc.add_table(rows=1, cols=4)
    kpis_table.style = 'Table Grid'
    
    hdr_cells = kpis_table.rows[0].cells
    hdr_cells[0].text = 'KPI Description'
    hdr_cells[1].text = 'Alignment to Existing KPI'
    hdr_cells[2].text = 'How is it Measured?'
    hdr_cells[3].text = 'Easily Measurable?'
    
    # Header Styling
    for cell in hdr_cells:
        set_cell_background(cell, "4F46E5") # Indigo
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                
    for idx, k in enumerate(st.session_state.kpis):
        row_cells = kpis_table.add_row().cells
        row_cells[0].text = k["kpi"].strip() if k["kpi"].strip() else "N/A"
        row_cells[1].text = k["alignment"].strip() if k["alignment"].strip() else "N/A"
        row_cells[2].text = k["measured"].strip() if k["measured"].strip() else "N/A"
        row_cells[3].text = k["easily_measurable"].strip() if k["easily_measurable"].strip() else "N/A"
        
        # Color alternate rows
        if idx % 2 == 1:
            for cell in row_cells:
                set_cell_background(cell, "F8FAFC")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Section 3: Business Impact
    h3 = doc.add_heading(level=1)
    r3 = h3.add_run("3. Business Impact Classifications")
    r3.font.color.rgb = RGBColor(79, 70, 229)
    
    impacts_table = doc.add_table(rows=1, cols=2)
    impacts_table.style = 'Table Grid'
    
    hdr_cells_i = impacts_table.rows[0].cells
    hdr_cells_i[0].text = 'Strategic Outcome Dimension'
    hdr_cells_i[1].text = 'Classification Level'
    
    for cell in hdr_cells_i:
        set_cell_background(cell, "4F46E5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                
    dimensions = [
        ("Revenue Influence", st.session_state.impact_revenue),
        ("Cost Reduction", st.session_state.impact_cost),
        ("Customer Experience Audience", st.session_state.impact_cust_service),
        ("Process Efficiency Outcome", st.session_state.impact_efficiency),
        ("Process Duration Outcome", st.session_state.impact_duration),
        ("Quality Improvement Domain", st.session_state.impact_quality),
    ]
    
    for idx, (dim, val) in enumerate(dimensions):
        row_cells = impacts_table.add_row().cells
        row_cells[0].text = dim
        row_cells[1].text = val
        if idx % 2 == 1:
            for cell in row_cells:
                set_cell_background(cell, "F8FAFC")
                
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Section 4: Sign-off
    h4 = doc.add_heading(level=1)
    r4 = h4.add_run("4. Authorizations & Approvals")
    r4.font.color.rgb = RGBColor(79, 70, 229)
    
    p_auth = doc.add_paragraph()
    p_auth.add_run("By signing below, the project sponsor and key stakeholders authorize allocation of resources to proceed with the planning and execution phases of this initiative.\n\n\n").font.color.rgb = RGBColor(71, 85, 105)
    
    p_sigs = doc.add_paragraph()
    p_sigs.add_run(
        "_____________________________________\t\t\t_____________________________________\n"
        "Project Sponsor Name & Date\t\t\tProject Lead Name & Date\n\n\n"
        "_____________________________________\t\t\t_____________________________________\n"
        "Finance Controller Name & Date\t\t\tIT Delivery Director Name & Date"
    ).font.size = Pt(10)
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Gemini Client caching resource
@st.cache_resource
def get_gemini_client(api_key):
    return genai.Client(api_key=api_key)

# SIDEBAR: Summary and Exports
with st.sidebar:
    st.markdown('<div style="text-align: center; margin-bottom: 20px;"><h2 style="color:#0f172a; font-size:1.5rem; font-weight:600;">📊 Console Panel</h2></div>', unsafe_allow_html=True)
    
    # Setup actions
    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.subheader("💡 Quick Setup")
    col_l, col_r = st.columns(2)
    with col_l:
        st.button("Load Demo", on_click=load_sample_data, use_container_width=True, type="primary")
    with col_r:
        st.button("Clear Form", on_click=reset_form, use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Progress Section
    filled_b, total_b, progress_pct = get_validation_stats()
    
    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.subheader("📈 Completion Status")
    
    st.markdown(f"""
    <div class="progress-label">
        <span>Required Fields Filled</span>
        <span>{progress_pct}%</span>
    </div>
    """, unsafe_allow_html=True)
    st.progress(progress_pct / 100.0)
    
    if progress_pct == 100:
        st.markdown('<div style="text-align:center; margin-top:10px;"><span class="status-badge badge-complete">Ready for Submission</span></div>', unsafe_allow_html=True)
    else:
        st.markdown('<div style="text-align:center; margin-top:10px;"><span class="status-badge badge-incomplete">Draft - Inputs Missing</span></div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # Form Live Summary Box
    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.subheader("📋 Executive Preview")
    
    proj_name = st.session_state.project_name if st.session_state.project_name.strip() else "*Unnamed Project*"
    st.markdown(f"**Project:** `{proj_name}`")
    
    st.markdown("**Business Impact Summary:**")
    st.markdown(f"- **Revenue:** `{st.session_state.impact_revenue}`")
    st.markdown(f"- **Cost:** `{st.session_state.impact_cost}`")
    st.markdown(f"- **Customer Service:** `{st.session_state.impact_cust_service}`")
    st.markdown(f"- **Process Efficiency:** `{st.session_state.impact_efficiency}`")
    st.markdown(f"- **Process Duration:** `{st.session_state.impact_duration}`")
    st.markdown(f"- **Quality:** `{st.session_state.impact_quality}`")
    
    st.markdown(f"**KPIs Defined:** `{len(st.session_state.kpis)}`")
    st.markdown('</div>', unsafe_allow_html=True)

    # INTEGRATION SETTINGS expander
    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.subheader("⚙️ Integration Settings")
    
    with st.expander("🔐 Databricks Warehouse", expanded=False):
        st.text_input("Server Hostname", key="db_host", help="Workspace Server Hostname (excluding https://)")
        st.text_input("HTTP Path", key="db_path", help="Warehouse http path found in Connection Details")
        st.text_input("Personal Access Token", key="db_token", type="password", help="Leave blank if deploying as Databricks App (uses OAuth)")
        st.markdown("---")
        st.text_input("Catalog", key="db_catalog")
        st.text_input("Schema", key="db_schema")
        st.text_input("Table Name", key="db_table")
        
    with st.expander("✨ Gemini API Config", expanded=False):
        st.text_input(
            "Gemini API Key",
            key="gemini_api_key",
            type="password",
            help="Get an API key from Google AI Studio. Required for the AI Copilot tab."
        )
    st.markdown('</div>', unsafe_allow_html=True)

    # Database Save Action
    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.subheader("📤 Database Action")
    
    if progress_pct < 100:
        st.warning("⚠️ Form is incomplete.")
        sidebar_submit_disabled = True
    else:
        st.success("✅ Form is ready to save.")
        sidebar_submit_disabled = False
        
    save_btn = st.button(
        "💾 Save to Databricks",
        use_container_width=True,
        type="primary",
        disabled=sidebar_submit_disabled,
        key="sidebar_save_btn"
    )
    
    if save_btn:
        payload = {
            "project_name": st.session_state.project_name,
            "problem_solving": st.session_state.problem_solving,
            "who_impacted": st.session_state.who_impacted,
            "cost_impact": st.session_state.cost_impact,
            "other_details": st.session_state.other_details,
            "value_creation": st.session_state.value_creation,
            "kpis": st.session_state.kpis,
            "impact_revenue": st.session_state.impact_revenue,
            "impact_cost": st.session_state.impact_cost,
            "impact_cust_service": st.session_state.impact_cust_service,
            "impact_efficiency": st.session_state.impact_efficiency,
            "impact_duration": st.session_state.impact_duration,
            "impact_quality": st.session_state.impact_quality
        }
        initiate_priority_and_save(payload)
    st.markdown('</div>', unsafe_allow_html=True)
    
    # Export Actions
    st.markdown('<div class="sidebar-section">', unsafe_allow_html=True)
    st.subheader("💾 Export Business Case")
    
    # Data Compilation
    sanitize_empty_fields()
    form_data = {
        "problem_statement": {
            "project_name": st.session_state.project_name,
            "problem_solving": st.session_state.problem_solving,
            "who_impacted": st.session_state.who_impacted,
            "cost_of_not_solving": st.session_state.cost_impact,
            "other_details": st.session_state.other_details
        },
        "value_kpis": {
            "value_creation_linkage": st.session_state.value_creation,
            "key_performance_indicators": st.session_state.kpis
        },
        "business_impact": {
            "revenue": st.session_state.impact_revenue,
            "cost": st.session_state.impact_cost,
            "customer_service": st.session_state.impact_cust_service,
            "process_efficiency": st.session_state.impact_efficiency,
            "process_duration": st.session_state.impact_duration,
            "quality": st.session_state.impact_quality
        }
    }
    
    json_str = json.dumps(form_data, indent=4)
    st.download_button(
        label="Download JSON",
        data=json_str,
        file_name=f"business_case_{st.session_state.project_name.lower().replace(' ', '_') or 'draft'}.json",
        mime="application/json",
        use_container_width=True
    )
    
    # Flatten Data to CSV for spreadsheet downloads
    flat_data = {
        "Project Name": [st.session_state.project_name],
        "Problem Solved": [st.session_state.problem_solving],
        "Who is Impacted": [st.session_state.who_impacted],
        "Cost of Not Solving": [st.session_state.cost_impact],
        "Other Details": [st.session_state.other_details],
        "Value Relation": [st.session_state.value_creation],
        "KPI Count": [len(st.session_state.kpis)],
        "Revenue Impact": [st.session_state.impact_revenue],
        "Cost Impact": [st.session_state.impact_cost],
        "Customer Service Impact": [st.session_state.impact_cust_service],
        "Process Efficiency Impact": [st.session_state.impact_efficiency],
        "Process Duration Impact": [st.session_state.impact_duration],
        "Quality Impact": [st.session_state.impact_quality]
    }
    df = pd.DataFrame(flat_data)
    csv_buffer = BytesIO()
    df.to_csv(csv_buffer, index=False)
    
    st.download_button(
        label="Download CSV",
        data=csv_buffer.getvalue(),
        file_name=f"business_case_{st.session_state.project_name.lower().replace(' ', '_') or 'draft'}.csv",
        mime="text/csv",
        use_container_width=True
    )
    
    # Word Charter docx export in sidebar
    if HAS_DOCX:
        doc_buffer = generate_project_charter_docx()
        if doc_buffer:
            st.download_button(
                label="📄 Download Word Charter",
                data=doc_buffer,
                file_name=f"project_charter_{st.session_state.project_name.lower().replace(' ', '_') or 'draft'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
    st.markdown('</div>', unsafe_allow_html=True)


# ==========================================
# PRIORITY AGENT HUMAN-IN-THE-LOOP INTERFACE
# ==========================================
if st.session_state.pending_priority_confirmation is not None:
    conf = st.session_state.pending_priority_confirmation
    calc_priority = conf["calculated_priority"]
    
    st.markdown(f"""
    <div style="background-color: #fef3c7; border-left: 5px solid #d97706; padding: 20px; border-radius: 8px; margin-bottom: 25px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05);">
        <h3 style="color: #92400e; margin: 0 0 10px 0; font-weight: 700;">⚡ Priority Agent Review Required</h3>
        <p style="color: #78350f; font-size: 16px; margin: 0 0 15px 0; line-height: 1.5;">
            The Priority Agent has calculated a priority assignment of <b>Priority {calc_priority}</b> for project <b>"{st.session_state.project_name or 'Draft'}"</b>. 
            Please review, choose whether to keep it or override, and click confirm to save to Databricks.
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    override_choice = st.radio(
        "**Confirm Priority Assignment:**",
        options=[
            f"Keep Calculated Priority ({calc_priority})",
            "Override to A (High revenue Increase & Saving)",
            "Override to B (External customer experience or Product quality)",
            "Override to C (Standard priority)"
        ],
        index=0,
        key="priority_choice_radio_main"
    )
    
    col_c1, col_c2 = st.columns(2)
    with col_c1:
        if st.button("✅ Confirm and Save to Databricks", type="primary", use_container_width=True, key="priority_confirm_btn"):
            if "Override to A" in override_choice:
                user_response = "A"
            elif "Override to B" in override_choice:
                user_response = "B"
            elif "Override to C" in override_choice:
                user_response = "C"
            else:
                user_response = calc_priority
                
            with st.spinner("Resuming Priority Agent workflow & saving to Databricks..."):
                final_event = run_priority_workflow_sync(
                    project_data=None,
                    user_response=user_response,
                    invocation_id=conf["invocation_id"]
                )
                
                if final_event and final_event.output and "priority" in final_event.output:
                    final_priority = final_event.output["priority"]
                    payload = conf["payload"]
                    payload["priority"] = final_priority
                    
                    success = save_to_databricks(payload)
                    if success:
                        st.balloons()
                        st.success(f"🎉 Saved successfully with Priority **{final_priority}** to Databricks!")
                        st.session_state.pending_priority_confirmation = None
                        st.session_state.priority_run_session_id = None
                        st.rerun()
                else:
                    st.error("❌ Priority Agent failed to return a valid output priority.")
                    
    with col_c2:
        if st.button("❌ Cancel Submission", use_container_width=True, key="priority_cancel_btn"):
            st.session_state.pending_priority_confirmation = None
            st.session_state.priority_run_session_id = None
            st.rerun()
            
    st.divider()

# TABS INITIALIZATION
tab_problem, tab_kpi, tab_impact, tab_copilot = st.tabs([
    "📝 Problem Statement", 
    "🎯 Value & KPIs", 
    "📊 Business Impact",
    "💬 AI Copilot"
])

# ==========================================
# TAB 1: PROBLEM STATEMENT
# ==========================================
with tab_problem:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("Project Identification & Problem Framing")
    st.write("Outline the project name, identify stakeholders, define the immediate pain point, and evaluate the status-quo cost.")
    
    # Project Name
    st.session_state.project_name = st.text_input(
        "PROJECT NAME",
        value=st.session_state.project_name,
        placeholder="e.g. Smart Order Automation Platform",
        help="Specify the unique identifying title of this project."
    )
    
    # Problem Description
    st.session_state.problem_solving = st.text_area(
        "WHAT SPECIFIC PROBLEM ARE WE SOLVING?",
        value=st.session_state.problem_solving,
        height=150,
        placeholder="Describe the challenges, roadblocks, and inefficiencies currently experienced...",
        help="Clearly articulate the core root cause problem here."
    )
    
    # Stakeholders Impacted
    st.session_state.who_impacted = st.text_area(
        "WHO IS IMPACTED?",
        value=st.session_state.who_impacted,
        height=100,
        placeholder="Specify the teams, customers, or systems affected by this problem...",
        help="Identify downstream departments, clients, or business units."
    )
    
    # Cost & Impact
    st.session_state.cost_impact = st.text_area(
        "WHAT IS THE CURRENT COST/IMPACT OF NOT SOLVING THE PROBLEM?",
        value=st.session_state.cost_impact,
        height=100,
        placeholder="Quantify the financial losses, time waste, compliance risks, or resource drain...",
        help="Specify financial values, hours lost, or qualitative risks."
    )
    
    # Other Details
    st.session_state.other_details = st.text_area(
        "OTHER DETAILS",
        value=st.session_state.other_details,
        height=100,
        placeholder="Add context, assumptions, background information, or dependency factors here...",
        help="Any additional notes or references."
    )
    
    st.markdown('</div>', unsafe_allow_html=True)


# ==========================================
# TAB 2: KPI & VALUE CREATION
# ==========================================
with tab_kpi:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("Value Realization Strategy")
    st.write("Define how this initiative ties into organizational value driver models and articulate success targets.")
    
    st.session_state.value_creation = st.text_area(
        "EXPLAIN HOW THE PROJECT DIRECTLY RELATES TO VALUE CREATION",
        value=st.session_state.value_creation,
        height=150,
        placeholder="e.g. Reduces cycle times by 30%, which unlocks direct labor savings and improves customer retention...",
        help="Demonstrate the business return link (e.g. cost reduction, capability expansion, revenue acceleration)."
    )
    st.markdown('</div>', unsafe_allow_html=True)
    
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("Key Performance Indicators (KPIs)")
    st.write("Add one or more metrics to evaluate successful delivery and alignment.")
    
    # Render dynamic inputs for KPIs list
    for idx, kpi_item in enumerate(st.session_state.kpis):
        st.markdown(f"##### KPI Metric #{idx + 1}")
        
        # Sub columns for KPI fields
        col1, col2 = st.columns(2)
        with col1:
            st.session_state.kpis[idx]["kpi"] = st.text_input(
                f"KPI #{idx+1} Description",
                value=kpi_item["kpi"],
                placeholder="What metric is being tracked? (e.g., Average Order Delivery Time)",
                key=f"kpi_desc_{idx}"
            )
            st.session_state.kpis[idx]["alignment"] = st.text_input(
                f"Alignment to Existing KPI (Metric #{idx+1})",
                value=kpi_item["alignment"],
                placeholder="Which parent enterprise strategic goal does this align to?",
                key=f"kpi_align_{idx}"
            )
        with col2:
            st.session_state.kpis[idx]["measured"] = st.text_input(
                f"How is it Measured? (Metric #{idx+1})",
                value=kpi_item["measured"],
                placeholder="How is this data queried or tabulated? (e.g. system logs)",
                key=f"kpi_measure_{idx}"
            )
            st.session_state.kpis[idx]["easily_measurable"] = st.text_input(
                f"Easily Measurable? (Metric #{idx+1})",
                value=kpi_item["easily_measurable"],
                placeholder="Is data easily accessible or requires extensive setup? (Yes/No with explanation)",
                key=f"kpi_ease_{idx}"
            )
            
        # Add visual divider if not the last item
        if idx < len(st.session_state.kpis) - 1:
            st.markdown("---")
            
    # Add/Delete buttons
    col_btn1, col_btn2, _ = st.columns([1, 1, 3])
    with col_btn1:
        if st.button("➕ Add Another KPI", key="add_kpi_btn"):
            st.session_state.kpis.append({
                "kpi": "",
                "alignment": "",
                "measured": "",
                "easily_measurable": ""
            })
            st.rerun()
    with col_btn2:
        if len(st.session_state.kpis) > 1:
            if st.button("🗑️ Remove Last KPI", key="remove_kpi_btn"):
                st.session_state.kpis.pop()
                st.rerun()
                
    st.markdown('</div>', unsafe_allow_html=True)


# ==========================================
# TAB 3: BUSINESS IMPACT
# ==========================================
with tab_impact:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("Strategic Core Business Impact Classification")
    st.write("Categorize how the outcomes of this project map to corporate impact vectors. Choose the most appropriate outcome direction.")
    
    # 1. Revenue
    st.write("##### 💵 Revenue Influence")
    revenue_options = ["No Impact", "Protection", "Increase"]
    st.session_state.impact_revenue = st.radio(
        "Select the Revenue alignment vector for this project:",
        options=revenue_options,
        index=revenue_options.index(st.session_state.impact_revenue),
        horizontal=True,
        key="impact_rev_radio"
    )
    
    st.markdown("---")
    
    # 2. Cost
    st.write("##### 📉 Cost Reduction")
    cost_options = ["No Impact", "Avoidance", "Saving"]
    st.session_state.impact_cost = st.radio(
        "Select the Cost outcome category:",
        options=cost_options,
        index=cost_options.index(st.session_state.impact_cost),
        horizontal=True,
        key="impact_cost_radio"
    )
    
    st.markdown("---")
    
    # 3. Customer Service
    st.write("##### 🤝 Customer Experience")
    cs_options = ["No impact", "External", "Internal"]
    st.session_state.impact_cust_service = st.radio(
        "Select the primary customer audience impacted:",
        options=cs_options,
        index=cs_options.index(st.session_state.impact_cust_service),
        horizontal=True,
        key="impact_cs_radio"
    )
    
    st.markdown("---")
    
    # 4. Process Efficiency
    st.write("##### ⚡ Process Efficiency")
    pe_options = ["No impact", "Improvement"]
    st.session_state.impact_efficiency = st.radio(
        "Select efficiency outcome impact:",
        options=pe_options,
        index=pe_options.index(st.session_state.impact_efficiency),
        horizontal=True,
        key="impact_eff_radio"
    )
    
    st.markdown("---")
    
    # 5. Process Duration
    st.write("##### ⏱️ Process Duration")
    pd_options = ["No impact", "Reduction"]
    st.session_state.impact_duration = st.radio(
        "Select cycle duration impact:",
        options=pd_options,
        index=pd_options.index(st.session_state.impact_duration),
        horizontal=True,
        key="impact_dur_radio"
    )
    
    st.markdown("---")
    
    # 6. Quality
    st.write("##### 🛡️ Quality Assurance")
    q_options = ["No impact", "Data", "Product"]
    st.session_state.impact_quality = st.radio(
        "Select the Quality improvement target domain:",
        options=q_options,
        index=q_options.index(st.session_state.impact_quality),
        horizontal=True,
        key="impact_q_radio"
    )
    
    st.markdown('</div>', unsafe_allow_html=True)


# ==========================================
# TAB 4: AI COPILOT / CHAT ASSISTANT
# ==========================================
with tab_copilot:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("💬 AI Copilot Business Case Assistant")
    st.write("Get clarification on the intake fields, ask how to improve your drafts, or brainstorm KPIs. The Copilot is restricted to this application's context.")
    
    # Check key
    api_key_to_use = st.session_state.gemini_api_key.strip()
    
    if not HAS_GEMINI:
        st.error("❌ 'google-genai' package is not installed. Run 'pip install google-genai' first.")
    elif not api_key_to_use:
        st.info("👋 To start chatting, please provide a **Gemini API Key** in the sidebar's **⚙️ Integration Settings** panel.", icon="🔑")
        
        # Also provide a convenience input inside the tab
        direct_key = st.text_input("Or enter Gemini API Key directly here:", type="password", key="direct_gemini_key")
        if direct_key:
            st.session_state.gemini_api_key = direct_key
            st.rerun()
    else:
        # Initialize Gemini Client
        try:
            client = get_gemini_client(api_key_to_use)
            
            # Formulate Strict System Prompt containing live field states
            system_instruction = f"""You are the Business Case Intake Copilot, an AI assistant strictly designed to help the user fill out the business case form on this page.
Your scope is STRICTLY restricted to:
1. Clarifying what each field in this portal means and how to fill it out:
   - PROJECT NAME (free text)
   - WHAT SPECIFIC PROBLEM ARE WE SOLVING? (free text description)
   - WHO IS IMPACTED? (stakeholders and business units affected)
   - WHAT IS THE CURRENT COST/IMPACT OF NOT SOLVING THE PROBLEM? (qualitative or financial costs)
   - OTHER DETAILS (context, assumptions)
   - EXPLAIN HOW THE PROJECT DIRECTLY RELATES TO VALUE CREATION (Narrative linkage)
   - KPI description, Alignment to existing KPI, How is it Measured?, Easily Measurable?
   - Business Impact dimensions (Revenue, Cost, Customer Service, Process Efficiency, Process Duration, Quality).
2. Reviewing the user's current draft answers and giving feedback/improvement recommendations.
3. Helping the user brainstorm descriptions, calculations, or KPI definitions for their project based on these fields.

DO NOT answer any general knowledge questions, write code, or help with topics unrelated to this business case intake form. If a user asks something out of scope, politely decline and steer them back to completing their business case. Keep your answers concise, structured, and helpful.

Here is the current live state of the user's form:
- Project Name: "{st.session_state.project_name}"
- Problem to Solve: "{st.session_state.problem_solving}"
- Who is Impacted: "{st.session_state.who_impacted}"
- Current Cost/Impact: "{st.session_state.cost_impact}"
- Other Details: "{st.session_state.other_details}"
- Value Creation Linkage: "{st.session_state.value_creation}"
- KPIs Defined: {json.dumps(st.session_state.kpis)}
- Business Impacts:
  - Revenue Influence: "{st.session_state.impact_revenue}"
  - Cost Reduction: "{st.session_state.impact_cost}"
  - Customer Experience: "{st.session_state.impact_cust_service}"
  - Process Efficiency: "{st.session_state.impact_efficiency}"
  - Process Duration: "{st.session_state.impact_duration}"
  - Quality: "{st.session_state.impact_quality}"
"""
            
            # Action: Suggested Prompt Shortcuts
            st.write("💡 **Quick Actions:**")
            col_pr1, col_pr2, col_pr3 = st.columns(3)
            
            suggested_prompt = None
            with col_pr1:
                if st.button("🔍 Review my current form draft", use_container_width=True):
                    suggested_prompt = "Can you review my current draft and suggest areas of improvement or missing context?"
            with col_pr2:
                if st.button("💡 Brainstorm KPIs for this project", use_container_width=True):
                    suggested_prompt = f"Based on my project '{st.session_state.project_name or 'Unnamed Project'}', can you suggest 2-3 specific KPIs I can add?"
            with col_pr3:
                if st.button("📊 Explain Business Impact options", use_container_width=True):
                    suggested_prompt = "What is the difference between Revenue Protection vs Increase, and Cost Avoidance vs Saving?"
                    
            # Clear chat button
            if st.button("🗑️ Clear Chat History"):
                st.session_state.copilot_messages = []
                st.rerun()
                
            st.markdown("---")
            
            # Render message logs
            chat_container = st.container(height=350)
            with chat_container:
                if not st.session_state.copilot_messages:
                    st.chat_message("assistant", avatar="✨").write("Hello! I am your Intake Copilot. Ask me how to define your project, clarify any form fields, or review your current draft.")
                else:
                    for msg in st.session_state.copilot_messages:
                        avatar = "👤" if msg["role"] == "user" else "✨"
                        st.chat_message(msg["role"], avatar=avatar).write(msg["content"])
            
            # Handle user text inputs or button triggers
            user_input = st.chat_input("Ask a question about the form fields...")
            
            # If a shortcut was clicked, prioritize it
            if suggested_prompt:
                user_input = suggested_prompt
                
            if user_input:
                # Add user message
                st.session_state.copilot_messages.append({"role": "user", "content": user_input})
                with chat_container:
                    st.chat_message("user", avatar="👤").write(user_input)
                    
                # Call Gemini API
                with chat_container:
                    with st.chat_message("assistant", avatar="✨"):
                        placeholder = st.empty()
                        full_reply = ""
                        
                        contents = []
                        for msg in st.session_state.copilot_messages:
                            role = "user" if msg["role"] == "user" else "model"
                            contents.append(
                                types.Content(
                                    role=role,
                                    parts=[types.Part.from_text(text=msg["content"])]
                                )
                            )
                            
                        try:
                            stream = client.models.generate_content_stream(
                                model="gemini-2.5-flash",
                                contents=contents,
                                config=types.GenerateContentConfig(
                                    system_instruction=system_instruction,
                                    temperature=0.7
                                )
                            )
                            for chunk in stream:
                                if chunk.text:
                                    full_reply += chunk.text
                                    placeholder.markdown(full_reply + "▌")
                            placeholder.markdown(full_reply)
                            
                            st.session_state.copilot_messages.append({"role": "assistant", "content": full_reply})
                        except Exception as e_gen:
                            st.error(f"Error calling Gemini: {str(e_gen)}")
                st.rerun()
                
        except Exception as e_init:
            st.error(f"Failed to initialize Gemini Client: {str(e_init)}")
            
    st.markdown('</div>', unsafe_allow_html=True)


# ==========================================
# CENTRAL SUBMISSION & DOCUMENT EXPORT SECTION
# ==========================================
st.markdown("---")
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.subheader("🚀 Central Databricks Submission & Documents")
st.write("Persist this business case record to your Databricks SQL Warehouse or download it as a Microsoft Word project charter.")

col_sub1, col_sub2 = st.columns([1, 1])
with col_sub1:
    st.write("##### 📤 Databricks SQL Sync")
    filled_b, total_b, progress_pct = get_validation_stats()
    
    if progress_pct < 100:
        st.warning("⚠️ Some required fields are currently empty. Please review all tabs before submitting.")
        submit_disabled = True
    else:
        st.success("✅ All required fields completed! Ready to submit.")
        submit_disabled = False
        
    submit_btn = st.button(
        "Submit Business Case to Databricks",
        use_container_width=True,
        type="primary",
        disabled=submit_disabled
    )
    
    if submit_btn:
        # Compile payload
        payload = {
            "project_name": st.session_state.project_name,
            "problem_solving": st.session_state.problem_solving,
            "who_impacted": st.session_state.who_impacted,
            "cost_impact": st.session_state.cost_impact,
            "other_details": st.session_state.other_details,
            "value_creation": st.session_state.value_creation,
            "kpis": st.session_state.kpis,
            "impact_revenue": st.session_state.impact_revenue,
            "impact_cost": st.session_state.impact_cost,
            "impact_cust_service": st.session_state.impact_cust_service,
            "impact_efficiency": st.session_state.impact_efficiency,
            "impact_duration": st.session_state.impact_duration,
            "impact_quality": st.session_state.impact_quality
        }
        initiate_priority_and_save(payload)
    
with col_sub2:
    st.write("##### 📄 Export Documents")
    if not HAS_DOCX:
        st.warning("⚠️ Microsoft Word exporter is unavailable. Run 'pip install python-docx'.")
    else:
        st.info("Generate and download a pre-formatted, fully editable Project Charter Word Document (`.docx`) containing all inputs.")
        doc_buffer = generate_project_charter_docx()
        if doc_buffer:
            st.download_button(
                label="📄 Generate & Download Word Charter",
                data=doc_buffer,
                file_name=f"project_charter_{st.session_state.project_name.lower().replace(' ', '_') or 'draft'}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
st.markdown('</div>', unsafe_allow_html=True)
